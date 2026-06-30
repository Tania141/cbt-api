"""
АКТ СИСТЕМ REST API v2.1 - Template-based document generation + Claude AI
"""
import os, io, re, json, base64
from datetime import datetime
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import openpyxl
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Border, Side
from docx import Document
import dropbox
from dropbox.exceptions import ApiError
from dropbox.files import WriteMode
import anthropic
import psycopg2
from psycopg2.extras import RealDictCursor
import bcrypt
import jwt
from functools import wraps
from validation_akt15 import validate_akt15

app = Flask(__name__)
CORS(app, origins="*")

DATABASE_URL          = os.environ.get("DATABASE_URL", "")
JWT_SECRET            = os.environ.get("JWT_SECRET", "change-me-in-production")
JWT_EXPIRY_HOURS      = int(os.environ.get("JWT_EXPIRY_HOURS", "24"))

DROPBOX_TOKEN         = os.environ.get("DROPBOX_TOKEN", "")
DROPBOX_REFRESH_TOKEN = os.environ.get("DROPBOX_REFRESH_TOKEN", "")
DROPBOX_APP_KEY       = os.environ.get("DROPBOX_APP_KEY", "")
DROPBOX_APP_SECRET    = os.environ.get("DROPBOX_APP_SECRET", "")
DROPBOX_FOLDER        = os.environ.get("DROPBOX_FOLDER", "/AKT_Projects")
ANTHROPIC_API_KEY     = os.environ.get("ANTHROPIC_API_KEY", "")

LOCAL_TEMPLATES_DIR   = os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates")

TEMPLATE_FILES = {
    "protokol2":         "Protokol_2_Template.docx",
    "protokol2combined": "Protokol_2_Combined_Template.docx",
    "protokol2p2":       "Protokol_2_Part2_Template.docx",
    "protokol2a":        "Protokol_2a_Template.docx",
    "obrazec3":    "template_obrazec3.docx",
    "akt5":        "Akt_5_Template.docx",
    "akt6":        "Akt_6_Template.docx",
    "akt7":        "Akt_7_Template.docx",
    "akt8":        "Akt_8_Template.docx",
    "akt9":        "Akt_9_Beton_Dnevnik_Template.docx",
    "akt10":       "Akt_10_Template.docx",
    "akt11":       "Akt_11_Template.docx",
    "akt13":       "Akt_13_Template.docx",
    "akt14":       "template_akt14_1.docx",
    "zapovedna":   "Zapovedna_Template.docx",
    "akt15":       "Akt_15_Template.docx",
    "akt16":       "Akt_16_Template.docx",
    "doklad":      "Okonchatelen_Doklad_Template.docx",
}

DOC_LABELS = {
    "protokol2":         "Протокол_2",
    "protokol2combined": "Protokol_2_Combined",
    "protokol2p2":       "Протокол_2_Част_2",
    "protokol2a":        "Протокол_2а",
    "obrazec3":    "Obrazec_3",
    "akt5":        "Akt_5",
    "akt6":        "Akt_6",
    "akt7":        "Akt_7",
    "akt8":        "Akt_8",
    "akt9":        "Akt_9_Beton_Dnevnik",
    "akt10":       "Akt_10",
    "akt11":       "Akt_11",
    "akt13":       "Akt_13",
    "akt14":       "Akt_14",
    "zapovedna":   "Zapovedna_Kniga",
    "akt15":       "Akt_15",
    "akt16":       "Akt_16",
    "doklad":      "Okonchatelen_Doklad",
}

# ── PostgreSQL ────────────────────────────────────────────────────────────────
def get_db():
    if not DATABASE_URL:
        return None
    return psycopg2.connect(DATABASE_URL, cursor_factory=RealDictCursor)

def init_db():
    conn = get_db()
    if not conn:
        print("DATABASE_URL не е зададен — PostgreSQL пропуснат")
        return
    with conn:
        with conn.cursor() as cur:
            cur.execute("""
                CREATE TABLE IF NOT EXISTS tenants (
                    id         SERIAL PRIMARY KEY,
                    name       TEXT NOT NULL,
                    email      TEXT NOT NULL UNIQUE,
                    created_at TIMESTAMPTZ DEFAULT NOW(),
                    is_active  BOOLEAN NOT NULL DEFAULT TRUE
                )
            """)
            cur.execute("""
                CREATE TABLE IF NOT EXISTS projects (
                    id          SERIAL PRIMARY KEY,
                    pi          TEXT NOT NULL,
                    tenant_id   TEXT NOT NULL,
                    stroej      TEXT,
                    address     TEXT,
                    consultant  TEXT,
                    passport    JSONB,
                    created_at  TIMESTAMPTZ DEFAULT NOW(),
                    updated_at  TIMESTAMPTZ DEFAULT NOW(),
                    UNIQUE (pi, tenant_id)
                )
            """)
            cur.execute("""
                CREATE TABLE IF NOT EXISTS audit_log (
                    id          SERIAL PRIMARY KEY,
                    action      TEXT NOT NULL,
                    user_id     TEXT,
                    tenant_id   TEXT,
                    endpoint    TEXT,
                    method      TEXT,
                    ip_address  TEXT,
                    user_agent  TEXT,
                    detail      JSONB,
                    created_at  TIMESTAMPTZ DEFAULT NOW()
                )
            """)
            # Migrate projects table — add columns introduced after initial deploy
            for col, typedef in [
                ("tenant_id",  "TEXT NOT NULL DEFAULT '1'"),
                ("consultant", "TEXT"),
                ("passport",   "JSONB"),
                ("updated_at", "TIMESTAMPTZ DEFAULT NOW()"),
            ]:
                cur.execute(f"ALTER TABLE projects ADD COLUMN IF NOT EXISTS {col} {typedef}")
            # Add UNIQUE constraint if missing (idempotent via DO NOTHING)
            cur.execute("""
                DO $$ BEGIN
                    IF NOT EXISTS (
                        SELECT 1 FROM pg_constraint
                        WHERE conname = 'projects_pi_tenant_id_key'
                    ) THEN
                        ALTER TABLE projects ADD CONSTRAINT projects_pi_tenant_id_key UNIQUE (pi, tenant_id);
                    END IF;
                END $$
            """)
            # Migrate existing audit_log tables that have the old schema
            for col, typedef in [
                ("user_id",    "TEXT"),
                ("tenant_id",  "TEXT"),
                ("endpoint",   "TEXT"),
                ("method",     "TEXT"),
                ("ip_address", "TEXT"),
                ("user_agent", "TEXT"),
                ("detail",     "JSONB"),
            ]:
                cur.execute(f"ALTER TABLE audit_log ADD COLUMN IF NOT EXISTS {col} {typedef}")
            cur.execute("""
                CREATE TABLE IF NOT EXISTS users (
                    id            SERIAL PRIMARY KEY,
                    email         TEXT NOT NULL UNIQUE,
                    password_hash TEXT NOT NULL,
                    role          TEXT NOT NULL DEFAULT 'viewer'
                                  CHECK (role IN ('admin', 'operator', 'viewer')),
                    tenant_id     TEXT,
                    created_at    TIMESTAMPTZ DEFAULT NOW(),
                    is_active     BOOLEAN NOT NULL DEFAULT TRUE
                )
            """)
    conn.close()
    print("PostgreSQL: таблиците са готови")

def log_action(action, user_id=None, tenant_id=None, detail=None):
    conn = get_db()
    if not conn:
        return
    try:
        with conn:
            with conn.cursor() as cur:
                cur.execute(
                    """INSERT INTO audit_log
                       (action, user_id, tenant_id, endpoint, method, ip_address, user_agent, detail)
                       VALUES (%s, %s, %s, %s, %s, %s, %s, %s)""",
                    (
                        action,
                        str(user_id) if user_id else None,
                        str(tenant_id) if tenant_id else None,
                        request.path,
                        request.method,
                        request.headers.get("X-Forwarded-For", request.remote_addr),
                        request.headers.get("User-Agent", ""),
                        json.dumps(detail, ensure_ascii=False) if detail else None,
                    )
                )
    except Exception:
        pass  # audit failures must never break the main request
    finally:
        conn.close()

# ── Dropbox helpers ───────────────────────────────────────────────────────────
def get_dropbox():
    if DROPBOX_REFRESH_TOKEN and DROPBOX_APP_KEY and DROPBOX_APP_SECRET:
        return dropbox.Dropbox(
            oauth2_refresh_token=DROPBOX_REFRESH_TOKEN,
            app_key=DROPBOX_APP_KEY,
            app_secret=DROPBOX_APP_SECRET
        )
    elif DROPBOX_TOKEN:
        return dropbox.Dropbox(DROPBOX_TOKEN)
    return None

def dbx_upload(dbx, path, data):
    dbx.files_upload(data, path, mode=WriteMode.overwrite, autorename=False, mute=True)

def dbx_download(dbx, path):
    _, res = dbx.files_download(path)
    return res.content

def dbx_create_folder(dbx, path):
    try:
        dbx.files_create_folder_v2(path)
    except ApiError as e:
        err_str = str(e)
        if "conflict" not in err_str and "already_exists" not in err_str:
            raise

def dbx_list_folder(dbx, path):
    try:
        r = dbx.files_list_folder(path)
        entries = list(r.entries)
        while r.has_more:
            r = dbx.files_list_folder_continue(r.cursor)
            entries.extend(r.entries)
        return entries
    except ApiError:
        return []

def get_shared_link(dbx, path):
    try:
        links = dbx.sharing_list_shared_links(path=path, direct_only=True).links
        if links:
            return links[0].url.replace("?dl=0", "?dl=1")
        return dbx.sharing_create_shared_link_with_settings(path).url.replace("?dl=0", "?dl=1")
    except Exception:
        return ""

# ── Excel helpers ─────────────────────────────────────────────────────────────
def build_passport_excel(rows):
    wb = Workbook()
    ws = wb.active
    ws.title = "Паспорт"
    ws.column_dimensions["A"].width = 38
    ws.column_dimensions["B"].width = 65
    hf  = Font(name="Calibri", bold=True, size=11, color="FFFFFF")
    hfill = PatternFill("solid", fgColor="1E3A5F")
    df  = Font(name="Calibri", size=11)
    af  = PatternFill("solid", fgColor="EAF1F8")
    thin = Side(style="thin", color="D0DBE6")
    brd  = Border(left=thin, right=thin, top=thin, bottom=thin)
    ws["A1"] = "Поле"; ws["B1"] = "Стойност"
    ws["A1"].font = hf; ws["B1"].font = hf
    ws["A1"].fill = hfill; ws["B1"].fill = hfill
    for i, row in enumerate(rows, start=2):
        if isinstance(row, (list, tuple)) and len(row) >= 2:
            ws.cell(i, 1, str(row[0])); ws.cell(i, 2, str(row[1]) if row[1] else "")
            ws.cell(i, 1).font = df; ws.cell(i, 2).font = df
            ws.cell(i, 1).border = brd; ws.cell(i, 2).border = brd
            if i % 2 == 0:
                ws.cell(i, 1).fill = af; ws.cell(i, 2).fill = af
    buf = io.BytesIO(); wb.save(buf); return buf.getvalue()

def read_passport_excel(data):
    ws = openpyxl.load_workbook(io.BytesIO(data)).active
    return {str(r[0]): str(r[1]) if r[1] else "" for r in ws.iter_rows(min_row=2, values_only=True) if r[0]}

def rows_to_dict(rows):
    return {str(r[0]): str(r[1]) if r[1] else "" for r in rows if isinstance(r, (list, tuple)) and len(r) >= 2}

# ── Placeholder logic ─────────────────────────────────────────────────────────
def two_names(s):
    parts = (s or "").strip().split()
    return " ".join(parts[:2])

def one_and_three(s):
    parts = (s or "").strip().split()
    if len(parts) >= 3:
        return f"{parts[0]} {parts[2]}"
    return " ".join(parts)

def fmt_date(v):
    if not v: return ""
    v = v.strip()
    if "." in v: return v
    try: return datetime.fromisoformat(v[:10]).strftime("%d.%m.%Y")
    except: return v

def extract_employees(d):
    n = int(d.get("Служители_Брой", 0) or 0)
    return [{
        "name": d.get(f"Служител_{i}_Име",""),
        "specialization": d.get(f"Служител_{i}_Специализация",""),
        "title": d.get(f"Служител_{i}_Титла","инж.")
    } for i in range(1, n+1) if d.get(f"Служител_{i}_Име","")]

def extract_projectants(d):
    n = int(d.get("Проектанти_Брой", 0) or 0)
    return [{
        "name": d.get(f"Проектант_{i}_Име",""),
        "ppp": d.get(f"Проектант_{i}_ППП",""),
        "specialization": d.get(f"Проектант_{i}_Специализация",""),
        "title": d.get(f"Проектант_{i}_Титла","инж.")
    } for i in range(1, n+1) if d.get(f"Проектант_{i}_Име","")]

def by_spec(items, spec):
    return [e for e in items if e.get("specialization") == spec]

def build_vazlogitel_block(d):
    tip   = d.get("Възложител_Тип", "Фирма")
    firma = d.get("Възложител_Фирма", "")
    adres = d.get("Възложител_Адрес", "")
    if tip in ("Физическо лице", "ФЛ"):
        return firma + (f", {adres}" if adres else "")
    eik  = d.get("Възложител_ЕИК", "")
    pred = d.get("Възложител_Представител", "")
    parts = [firma]
    if eik:   parts.append(f"ЕИК {eik}")
    if adres: parts.append(adres)
    if pred:  parts.append(f"представлявано от {pred}")
    return ", ".join(parts)

def build_projectants_list(projectants):
    lines = []
    for p in projectants:
        spec = p["specialization"]
        title = p["title"]
        name = p["name"]
        ppp = p["ppp"]
        kamara = "КАБ" if spec in ("Архитектура", "Паркоустройство и Благоустройство") else "КИИП"
        line = "Част " + spec + ": " + title + " " + name
        if ppp:
            line += ", рег. № " + ppp + " в " + kamara
        lines.append(line)
    return "\n".join(lines)

def build_employees_list(employees):
    return "\n".join(f"Част {e['specialization']}: {e['title']} {e['name']}" for e in employees)

def build_projectants_signatures(projectants):
    return "\n".join(
        f"Част {p['specialization']}: {p['title']} {p['name']} ….................................................................."
        for p in projectants
    )

def build_employees_signatures(employees):
    return "\n".join(
        f"Част {e['specialization']}: {e['title']} {e['name']} ….................................................................."
        for e in employees
    )

def build_placeholders(d):
    employees   = extract_employees(d)
    projectants = extract_projectants(d)

    geo     = (by_spec(employees, "Геодезия") or [{}])[0].get("name", d.get("Геодезист",""))
    sn_k    = (by_spec(employees, "Конструктивна") or [{}])[0].get("name", "")
    pj_k    = (by_spec(projectants, "Конструктивна") or [{}])[0].get("name", "")
    pj_arch = (by_spec(projectants, "Архитектура") or [{}])[0].get("name", "")
    specs   = "; ".join(f"{e['title']} {e['name']} ({e['specialization']})" for e in employees) or d.get("Консултант_Управител","")
    vaz_tip = d.get("Възложител_Тип", "Фирма")
    upr     = d.get("Консултант_Управител", "")
    teh_ryk = d.get("Строител_ТехРък", "")
    str_upr = d.get("Строител_Управител", "")
    vaz_pr  = d.get("Възложител_Представител", "")
    # Възложител_1и3: prefer explicit signing person, fall back to ФЛ name or representative
    vaz_podpisva = d.get("Възложител_Подписва", "")
    vaz_name_for_1i3 = vaz_podpisva or (d.get("Възложител_Фирма", "") if vaz_tip in ("Физическо лице", "ФЛ") else vaz_pr)

    return {
        "{{Строеж}}":                    d.get("Строеж",""),
        "{{Адрес}}":                     d.get("Адрес",""),
        "{{Консултант_Фирма}}":          d.get("Консултант_Фирма",""),
        "{{Консултант_ЕИК}}":            d.get("Консултант_ЕИК",""),
        "{{Консултант_Адрес}}":          d.get("Консултант_Адрес",""),
        "{{Консултант_Управител}}":      upr,
        "{{Консултант_Удостоверение}}":   d.get("Консултант_Удостоверение",""),
        "{{Управител_2имена}}":          two_names(upr),
        "{{Управител_1и3}}":             one_and_three(upr),
        "{{Строител_Фирма}}":            d.get("Строител_Фирма",""),
        "{{Строител_ЕИК}}":              d.get("Строител_ЕИК",""),
        "{{Строител_Адрес}}":            d.get("Строител_Адрес",""),
        "{{Строител_Управител}}":        str_upr,
        "{{Строител_Управител_2имена}}": two_names(str_upr),
        "{{Строител_Управител_1и3}}":    one_and_three(str_upr),
        "{{Строител_ТехРък}}":           teh_ryk,
        "{{ТехРък_2имена}}":             two_names(teh_ryk),
        "{{ТехРък_1и3}}":               one_and_three(teh_ryk),
        "{{tech_director}}":             teh_ryk,
        "{{Възложител_Тип}}":            vaz_tip,
        "{{Възложител_Фирма}}":          d.get("Възложител_Фирма",""),
        "{{Възложител_ЕИК}}":            d.get("Възложител_ЕИК","") if vaz_tip not in ("Физическо лице","ФЛ") else "",
        "{{Възложител_Адрес}}":          d.get("Възложител_Адрес",""),
        "{{Възложител_Представител}}":   vaz_pr if vaz_tip not in ("Физическо лице","ФЛ") else "",
        "{{Възложител_2имена}}":         two_names(vaz_name_for_1i3),
        "{{Възложител_1и3}}":           one_and_three(vaz_name_for_1i3),
        "{{Възложител_Блок}}":           build_vazlogitel_block(d),   # и (correct, matches template)
        "{{Възложател_Блок}}":           build_vazlogitel_block(d),   # а (legacy alias)
        "{{Възложител_Подписва_Блок}}":  d.get("Възложител_Подписва", "") or build_vazlogitel_block(d),  # и
        "{{Възложател_Подписва_Блок}}":  d.get("Възложател_Подписва", "") or build_vazlogitel_block(d),  # а
        "{{РС_Номер}}":                  d.get("РС_Номер",""),
        "{{РС_Дата}}":                   fmt_date(d.get("РС_Дата","")),
        "{{РС_Издател}}":                d.get("РС_Издател",""),
        "{{РС_ВСила}}":                  fmt_date(d.get("РС_ВСила","")),
        "{{Геодезист}}":                 geo,
        "{{Геодезист_2имена}}":          two_names(geo),
        "{{Геодезист_1и3}}":            one_and_three(geo),
        "{{consultant_specialists}}":    specs,
        "{{constructor_name}}":          pj_k,
        "{{sn_konstruktivna}}":          sn_k,
        "{{СН_Конструктивна}}":          sn_k,
        "{{СН_Архитектура}}":            (by_spec(employees, "Архитектура") or [{}])[0].get("name", ""),
        "{{СН_Електро}}":                (by_spec(employees, "Електро") or [{}])[0].get("name", ""),
        "{{СН_ВиК}}":                    (by_spec(employees, "ВиК") or [{}])[0].get("name", ""),
        "{{СН_Геодезия}}":               (by_spec(employees, "Геодезия") or [{}])[0].get("name", ""),
        "{{СН_ПБ}}":                     (by_spec(employees, "ПБ") or [{}])[0].get("name", ""),
        "{{СН_Пътна}}":                  (by_spec(employees, "Пътна") or [{}])[0].get("name", ""),
        "{{СН_ОВК}}":                    (by_spec(employees, "ОВК и ЕЕ") or [{}])[0].get("name", ""),
        "{{pj_konstruktivna}}":          pj_k,
        "{{Конструктивна}}":             pj_k,
        "{{ПЖ_Конструктивна}}":          pj_k,
        "{{ПЖ_Архитектура}}":            pj_arch,
        "{{Вода}}":                      d.get("Вода", ""),
        "{{Канализация}}":               d.get("Канализация", ""),
        "{{Ел_Захранване}}":             d.get("Ел_Захранване", ""),
        "{{Проектанти_Списък}}":         build_projectants_list(projectants),
        "{{Консултанти_Списък}}":        build_employees_list(employees),
        "{{Проектанти_Подписи}}":        build_projectants_signatures(projectants),
        "{{Консултанти_Подписи}}":       build_employees_signatures(employees),
        "{{Opisanie_Ploshtadka}}":       d.get("Opisanie_Ploshtadka", ""),
        "{{Sastoyanie_Okolo}}":          d.get("Sastoyanie_Okolo", "пътните и тротоарни настилки по прилежащата улица са в добро състояние, съседните имоти няма да бъдат засягани от бъдещото строителство"),
        "{{Merki_PBZ}}":                 d.get("Merki_PBZ", "ще се осъществява от прилежащата улична мрежа съгласно съгласуван ПБЗ"),
        "{{Darvesenost}}":               d.get("Darvesenost", ""),
        "{{Kota_Izkop}}":               d.get("Kota_Izkop", ""),
        "{{Kota_Cokul}}":               d.get("Kota_Cokul", ""),
        "{{Kota_Korniz}}":              d.get("Kota_Korniz", ""),
        "{{Kota_Bilo}}":                d.get("Kota_Bilo", ""),
        "{{Reper_Nomer}}":              d.get("Reper_Nomer", ""),
        "{{Reper_Kota}}":               d.get("Reper_Kota", ""),
    }


# ── Template engine ───────────────────────────────────────────────────────────
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

def insert_paragraphs_after(para, lines, font_name="Times New Roman", font_size=12):
    from docx.shared import Pt
    ref = para._element
    parent = ref.getparent()
    idx = list(parent).index(ref)
    for i, line in enumerate(lines):
        new_p = OxmlElement("w:p")
        new_r = OxmlElement("w:r")
        new_rpr = OxmlElement("w:rPr")
        if para.runs:
            orig_rpr = para.runs[0]._r.find(qn("w:rPr"))
            if orig_rpr is not None:
                new_rpr = deepcopy(orig_rpr)
        new_r.append(new_rpr)
        new_t = OxmlElement("w:t")
        new_t.text = line
        new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        new_r.append(new_t)
        new_p.append(new_r)
        parent.insert(idx + 1 + i, new_p)

def replace_in_runs(para, replacements):
    full = "".join(r.text for r in para.runs)
    if not any(k in full for k in replacements):
        return False

    multiline_key = None
    multiline_val = None
    for k, v in replacements.items():
        if k in full and "\n" in v:
            multiline_key = k
            multiline_val = v
            break

    if multiline_key:
        lines = multiline_val.split("\n")
        first_line = full.replace(multiline_key, lines[0])
        for k, v in replacements.items():
            if k != multiline_key and "\n" not in v:
                first_line = first_line.replace(k, v)
        if para.runs:
            para.runs[0].text = first_line
            for r in para.runs[1:]:
                r.text = ""
        if len(lines) > 1:
            insert_paragraphs_after(para, lines[1:])
        return True

    new_text = full
    for k, v in replacements.items():
        new_text = new_text.replace(k, v)
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""
    return False

def fill_template(doc, replacements):
    paras = list(doc.paragraphs)
    for para in paras:
        replace_in_runs(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_runs(para, replacements)
    for section in doc.sections:
        for para in section.header.paragraphs:
            replace_in_runs(para, replacements)
        for para in section.footer.paragraphs:
            replace_in_runs(para, replacements)
    return doc

def generate_from_template(template_name, replacements):
    path = os.path.join(LOCAL_TEMPLATES_DIR, template_name)
    if not os.path.isfile(path):
        raise FileNotFoundError(f"Шаблонът '{template_name}' липсва локално → {path}")
    doc = Document(path)
    fill_template(doc, replacements)
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

# ── Auth helpers ─────────────────────────────────────────────────────────────
def make_token(user):
    from datetime import timezone, timedelta
    payload = {
        "sub":       str(user["id"]),
        "email":     user["email"],
        "role":      user["role"],
        "tenant_id": user["tenant_id"],
        "exp":       datetime.now(timezone.utc) + timedelta(hours=JWT_EXPIRY_HOURS),
    }
    return jwt.encode(payload, JWT_SECRET, algorithm="HS256")

def require_auth(f):
    @wraps(f)
    def wrapper(*args, **kwargs):
        auth = request.headers.get("Authorization", "")
        if not auth.startswith("Bearer "):
            return jsonify({"error": "Липсва Authorization header"}), 401
        try:
            payload = jwt.decode(auth[7:], JWT_SECRET, algorithms=["HS256"])
        except jwt.ExpiredSignatureError:
            return jsonify({"error": "Токенът е изтекъл"}), 401
        except jwt.InvalidTokenError:
            return jsonify({"error": "Невалиден токен"}), 401
        request.current_user = payload
        return f(*args, **kwargs)
    return wrapper

def require_admin(f):
    @wraps(f)
    @require_auth
    def wrapper(*args, **kwargs):
        if request.current_user.get("role") != "admin":
            return jsonify({"error": "Само администратори имат достъп"}), 403
        return f(*args, **kwargs)
    return wrapper

# ── Routes ────────────────────────────────────────────────────────────────────
@app.route("/health", methods=["GET"])
def health():
    dbx = get_dropbox()
    dbx_ok = False
    if dbx:
        try: dbx.users_get_current_account(); dbx_ok = True
        except: pass

    templates = {}
    for key, fname in TEMPLATE_FILES.items():
        local_path = os.path.join(LOCAL_TEMPLATES_DIR, fname)
        templates[key] = f"OK ({fname})" if os.path.isfile(local_path) else f"ЛИПСВА ({fname})"

    db_ok = False
    db_error = None
    try:
        conn = get_db()
        if conn:
            with conn.cursor() as cur:
                cur.execute("SELECT 1")
            conn.close()
            db_ok = True
    except Exception as e:
        db_error = str(e)

    return jsonify({
        "status": "ok", "version": "2.1",
        "timestamp": datetime.utcnow().isoformat(),
        "dropbox": "свързан" if dbx_ok else "не е свързан",
        "dropbox_folder": DROPBOX_FOLDER,
        "templates_dir": LOCAL_TEMPLATES_DIR,
        "templates": templates,
        "ai": "конфигуриран" if ANTHROPIC_API_KEY else "не е конфигуриран",
        "database": "свързан" if db_ok else ("не е конфигуриран" if not DATABASE_URL else f"грешка: {db_error}"),
        "build": "10af59c",
    })

@app.route("/api/auth/register", methods=["POST"])
def auth_register():
    body = request.get_json()
    if not body or not body.get("email") or not body.get("password"):
        return jsonify({"error": "Необходими са email и password"}), 400
    if not body.get("tenant_id"):
        return jsonify({"error": "Необходим е tenant_id"}), 400

    role      = body.get("role", "viewer")
    tenant_id = str(body["tenant_id"])
    if role not in ("admin", "operator", "viewer"):
        return jsonify({"error": "Невалидна роля. Позволени: admin, operator, viewer"}), 400

    pw_hash = bcrypt.hashpw(body["password"].encode(), bcrypt.gensalt()).decode()

    conn = get_db()
    if not conn:
        return jsonify({"error": "База данни не е свързана"}), 503
    try:
        with conn.cursor() as cur:
            cur.execute("SELECT id FROM tenants WHERE id = %s AND is_active = TRUE", (tenant_id,))
            if not cur.fetchone():
                return jsonify({"error": "Невалиден или неактивен tenant_id"}), 400
        with conn:
            with conn.cursor() as cur:
                cur.execute(
                    """INSERT INTO users (email, password_hash, role, tenant_id)
                       VALUES (%s, %s, %s, %s) RETURNING id, email, role, tenant_id, created_at""",
                    (body["email"].lower().strip(), pw_hash, role, tenant_id)
                )
                user = dict(cur.fetchone())
        return jsonify({"status": "ok", "user": user}), 201
    except psycopg2.errors.UniqueViolation:
        return jsonify({"error": "Email вече е регистриран"}), 409
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        conn.close()


@app.route("/api/auth/login", methods=["POST"])
def auth_login():
    body = request.get_json()
    if not body or not body.get("email") or not body.get("password"):
        return jsonify({"error": "Необходими са email и password"}), 400

    conn = get_db()
    if not conn:
        return jsonify({"error": "База данни не е свързана"}), 503
    try:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT * FROM users WHERE email = %s",
                (body["email"].lower().strip(),)
            )
            user = cur.fetchone()
        if not user or not bcrypt.checkpw(body["password"].encode(), user["password_hash"].encode()):
            return jsonify({"error": "Невалиден email или парола"}), 401
        if not user["is_active"]:
            return jsonify({"error": "Акаунтът е деактивиран"}), 403
        token = make_token(user)
        log_action("login", user_id=user["id"], tenant_id=user["tenant_id"],
                   detail={"email": user["email"], "role": user["role"]})
        return jsonify({"status": "ok", "token": token, "role": user["role"]})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        conn.close()


@app.route("/api/auth/me", methods=["GET"])
@require_auth
def auth_me():
    conn = get_db()
    if not conn:
        return jsonify({"error": "База данни не е свързана"}), 503
    try:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT id, email, role, tenant_id, created_at, is_active FROM users WHERE id = %s",
                (request.current_user["sub"],)
            )
            user = cur.fetchone()
        if not user:
            return jsonify({"error": "Потребителят не е намерен"}), 404
        log_action("get_me", user_id=request.current_user["sub"],
                   tenant_id=request.current_user.get("tenant_id"))
        return jsonify(dict(user))
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        conn.close()


@app.route("/api/setup/init", methods=["POST"])
def setup_init():
    body = request.get_json()
    required = ("tenant_name", "tenant_email", "admin_email", "admin_password")
    if not body or any(not body.get(k) for k in required):
        return jsonify({"error": f"Необходими полета: {', '.join(required)}"}), 400

    conn = get_db()
    if not conn:
        return jsonify({"error": "База данни не е свързана"}), 503
    try:
        with conn.cursor() as cur:
            cur.execute("SELECT COUNT(*) AS n FROM users")
            if cur.fetchone()["n"] > 0:
                return jsonify({"error": "Системата вече е инициализирана"}), 403

        with conn:
            with conn.cursor() as cur:
                cur.execute(
                    "INSERT INTO tenants (name, email) VALUES (%s, %s) RETURNING id, name, email",
                    (body["tenant_name"].strip(), body["tenant_email"].lower().strip())
                )
                tenant = dict(cur.fetchone())

                pw_hash = bcrypt.hashpw(body["admin_password"].encode(), bcrypt.gensalt()).decode()
                cur.execute(
                    """INSERT INTO users (email, password_hash, role, tenant_id)
                       VALUES (%s, %s, 'admin', %s) RETURNING id, email, role, tenant_id""",
                    (body["admin_email"].lower().strip(), pw_hash, str(tenant["id"]))
                )
                admin = dict(cur.fetchone())

        return jsonify({"status": "ok", "tenant": tenant, "admin": admin}), 201
    except psycopg2.errors.UniqueViolation:
        return jsonify({"error": "Tenant email или admin email вече съществуват"}), 409
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        conn.close()


@app.route("/api/admin/tenants", methods=["POST"])
@require_admin
def create_tenant():
    body = request.get_json()
    if not body or not body.get("name") or not body.get("email"):
        return jsonify({"error": "Необходими са name и email"}), 400
    conn = get_db()
    if not conn:
        return jsonify({"error": "База данни не е свързана"}), 503
    try:
        with conn:
            with conn.cursor() as cur:
                cur.execute(
                    """INSERT INTO tenants (name, email)
                       VALUES (%s, %s) RETURNING id, name, email, created_at, is_active""",
                    (body["name"].strip(), body["email"].lower().strip())
                )
                tenant = dict(cur.fetchone())
        return jsonify({"status": "ok", "tenant": tenant}), 201
    except psycopg2.errors.UniqueViolation:
        return jsonify({"error": "Tenant с този email вече съществува"}), 409
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        conn.close()


@app.route("/api/admin/tenants", methods=["GET"])
@require_admin
def list_tenants():
    conn = get_db()
    if not conn:
        return jsonify({"error": "База данни не е свързана"}), 503
    try:
        with conn.cursor() as cur:
            cur.execute("SELECT id, name, email, created_at, is_active FROM tenants ORDER BY created_at DESC")
            tenants = [dict(r) for r in cur.fetchall()]
        return jsonify({"tenants": tenants, "count": len(tenants)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        conn.close()


@app.route("/api/admin/audit-log", methods=["GET"])
@require_admin
def get_audit_log():
    conn = get_db()
    if not conn:
        return jsonify({"error": "База данни не е свързана"}), 503
    try:
        filters, params = [], []

        user_id   = request.args.get("user_id")
        action    = request.args.get("action")
        date_from = request.args.get("date_from")
        date_to   = request.args.get("date_to")

        if user_id:
            filters.append("user_id = %s");   params.append(user_id)
        if action:
            filters.append("action = %s");    params.append(action)
        if date_from:
            filters.append("created_at >= %s"); params.append(date_from)
        if date_to:
            filters.append("created_at <= %s"); params.append(date_to)

        where = ("WHERE " + " AND ".join(filters)) if filters else ""
        params.append(100)

        with conn.cursor() as cur:
            cur.execute(
                f"""SELECT id, action, user_id, tenant_id, endpoint, method,
                           ip_address, user_agent, detail, created_at
                    FROM audit_log {where}
                    ORDER BY created_at DESC
                    LIMIT %s""",
                params
            )
            entries = [dict(r) for r in cur.fetchall()]
        return jsonify({"entries": entries, "count": len(entries)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        conn.close()


@app.route("/api/passports", methods=["GET"])
@require_auth
def list_passports():
    try:
        tenant_id = request.current_user.get("tenant_id")
        if not tenant_id:
            return jsonify({"error": "Токенът не съдържа tenant_id"}), 403
        conn = get_db()
        if not conn:
            return jsonify({"error": "База данни не е конфигурирана"}), 503
        with conn.cursor() as cur:
            cur.execute(
                "SELECT pi, stroej, address, passport, created_at, updated_at "
                "FROM projects WHERE tenant_id = %s ORDER BY updated_at DESC",
                (str(tenant_id),)
            )
            rows = cur.fetchall()
        conn.close()
        result = []
        for r in rows:
            passport_data = r["passport"] or {}
            entry = dict(passport_data) if isinstance(passport_data, dict) else {}
            entry["pi"]        = r["pi"]
            entry["stroej"]    = r["stroej"] or ""
            entry["address"]   = r["address"] or ""
            # consultant stays as JSONB object {id, name} from passport — TEXT column not used
            entry["createdAt"] = r["created_at"].isoformat() if r["created_at"] else None
            entry["updatedAt"] = r["updated_at"].isoformat() if r["updated_at"] else None
            result.append(entry)
        log_action("get_passports", user_id=request.current_user["sub"], tenant_id=tenant_id,
                   detail={"count": len(result)})
        def _default(o):
            if isinstance(o, datetime): return o.isoformat()
            return str(o)
        from flask import Response as FlaskResponse
        return FlaskResponse(json.dumps(result, default=_default), mimetype="application/json")
    except Exception as e:
        import traceback
        return jsonify({"error": str(e), "type": type(e).__name__,
                        "trace": traceback.format_exc()}), 500

@app.route("/api/passports/<pi>", methods=["GET"])
@require_auth
def get_passport(pi):
    tenant_id = request.current_user.get("tenant_id")
    if not tenant_id:
        return jsonify({"error": "Токенът не съдържа tenant_id"}), 403
    dbx = get_dropbox()
    if not dbx: return jsonify({"error": "Dropbox не е конфигуриран"}), 503
    try:
        data = dbx_download(dbx, f"{DROPBOX_FOLDER}/{tenant_id}/PI-{pi}/passport.xlsx")
        log_action("get_passport", user_id=request.current_user["sub"], tenant_id=tenant_id,
                   detail={"pi": pi})
        return jsonify({"pi": pi, "tenant_id": tenant_id, "passport": read_passport_excel(data)})
    except ApiError:
        return jsonify({"error": f"Паспорт PI-{pi} не е намерен"}), 404
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/passports/<pi>", methods=["POST"])
@require_auth
def save_passport(pi):
    try:
        tenant_id = request.current_user.get("tenant_id")
        if not tenant_id:
            return jsonify({"error": "Токенът не съдържа tenant_id"}), 403
        body = request.get_json()
        if not body or "passport" not in body:
            return jsonify({"error": "Липсва поле 'passport'"}), 400
        passport = body["passport"]
        conn = get_db()
        if not conn:
            return jsonify({"error": "База данни не е конфигурирана"}), 503
        stroej     = passport.get("stroej", "") if isinstance(passport, dict) else ""
        address    = passport.get("address", "") if isinstance(passport, dict) else ""
        consultant = (passport.get("consultant") or {}).get("name", "") if isinstance(passport, dict) else ""
        with conn:
            with conn.cursor() as cur:
                cur.execute("""
                    INSERT INTO projects (pi, tenant_id, stroej, address, consultant, passport, updated_at)
                    VALUES (%s, %s, %s, %s, %s, %s::jsonb, NOW())
                    ON CONFLICT (pi, tenant_id) DO UPDATE
                        SET stroej = EXCLUDED.stroej, address = EXCLUDED.address,
                            consultant = EXCLUDED.consultant, passport = EXCLUDED.passport,
                            updated_at = NOW()
                """, (pi, str(tenant_id), stroej, address, consultant, json.dumps(passport)))
        conn.close()
        log_action("save_passport", user_id=request.current_user["sub"], tenant_id=tenant_id,
                   detail={"pi": pi})
        return jsonify({"status": "ok", "pi": pi, "tenant_id": tenant_id})
    except Exception as e:
        import traceback
        return jsonify({"error": str(e), "type": type(e).__name__,
                        "trace": traceback.format_exc()}), 500

@app.route("/api/generate/<doc_type>", methods=["POST"])
@require_auth
def generate_document(doc_type):
    tenant_id = request.current_user.get("tenant_id")
    if not tenant_id:
        return jsonify({"error": "Токенът не съдържа tenant_id"}), 403
    if doc_type not in TEMPLATE_FILES:
        return jsonify({"error": f"Непознат тип: {doc_type}. Позволени: {list(TEMPLATE_FILES.keys())}"}), 400
    body = request.get_json()
    if not body or "passport" not in body:
        return jsonify({"error": "Липсва поле 'passport'"}), 400

    d    = rows_to_dict(body["passport"])
    pi   = str(body.get("pi", "unknown"))
    repl = build_placeholders(d)

    if doc_type == "protokol2combined":
        print(f"[DEBUG protokol2combined] PI={pi} keys={sorted(d.keys())}", flush=True)
        print(f"[DEBUG protokol2combined] Геодезист={d.get('Геодезист','MISSING')} ПЖ_Конструктивна={repl.get('{{ПЖ_Конструктивна}}','MISSING')} ПЖ_Архитектура={repl.get('{{ПЖ_Архитектура}}','MISSING')}", flush=True)

    if doc_type == "zapovedna":
        repl["{{Заповедна_Номер}}"] = body.get("zapovedna_number", "___")
        repl["{{Заповедна_Дата}}"]  = fmt_date(body.get("zapovedna_date", ""))

    download_name = f"{doc_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    try:
        doc_bytes = generate_from_template(TEMPLATE_FILES[doc_type], repl)

        # Optional Dropbox upload — only if credentials are configured
        file_url = None
        if DROPBOX_REFRESH_TOKEN:
            dbx = get_dropbox()
            if dbx:
                pi_filename = f"PI-{pi}_{DOC_LABELS[doc_type]}.docx"
                folder = f"{DROPBOX_FOLDER}/{tenant_id}/PI-{pi}"
                path   = f"{folder}/{pi_filename}"
                try:
                    dbx_create_folder(dbx, folder)
                    dbx_upload(dbx, path, doc_bytes)
                    file_url = get_shared_link(dbx, path)
                except Exception:
                    pass  # Dropbox failure does not block the download

        log_action("generate_doc", user_id=request.current_user["sub"], tenant_id=tenant_id,
                   detail={"doc_type": doc_type, "pi": pi, "filename": download_name,
                           "dropbox_url": file_url})

        buf = io.BytesIO(doc_bytes)
        buf.seek(0)
        response = send_file(
            buf,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=download_name,
        )
        if file_url:
            response.headers["X-Dropbox-URL"] = file_url
        return response

    except FileNotFoundError as e:
        return jsonify({"error": str(e),
                        "hint": f"Постави шаблона в папка: {LOCAL_TEMPLATES_DIR}"}), 404
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ── Cloud Sync ────────────────────────────────────────────────────────────────
@app.route("/api/cloud/save", methods=["POST"])
@require_auth
def cloud_save():
    tenant_id = request.current_user.get("tenant_id")
    if not tenant_id:
        return jsonify({"error": "Токенът не съдържа tenant_id"}), 403
    dbx = get_dropbox()
    if not dbx:
        return jsonify({"error": "Dropbox не е свързан"}), 503
    body = request.get_json()
    if not body:
        return jsonify({"error": "Липсва тяло"}), 400
    path = f"{DROPBOX_FOLDER}/{tenant_id}/_akt_projects.json"
    try:
        data = json.dumps(body, ensure_ascii=False, indent=2).encode('utf-8')
        dbx_upload(dbx, path, data)
        log_action("cloud_save", user_id=request.current_user["sub"], tenant_id=tenant_id,
                   detail={"path": path})
        return jsonify({"status": "ok", "path": path})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/cloud/load", methods=["GET"])
@require_auth
def cloud_load():
    tenant_id = request.current_user.get("tenant_id")
    if not tenant_id:
        return jsonify({"error": "Токенът не съдържа tenant_id"}), 403
    dbx = get_dropbox()
    if not dbx:
        return jsonify({"error": "Dropbox не е свързан"}), 503
    path = f"{DROPBOX_FOLDER}/{tenant_id}/_akt_projects.json"
    try:
        data = dbx_download(dbx, path)
        payload = json.loads(data.decode('utf-8'))
        log_action("cloud_load", user_id=request.current_user["sub"], tenant_id=tenant_id)
        return jsonify(payload)
    except Exception:
        return jsonify({"projects": [], "consultants": [], "history": []}), 200


# ── AI Протоколи ──────────────────────────────────────────────────────────────
@app.route("/api/ai/generate", methods=["POST"])
def ai_generate():
    """
    Приема prompt + (опционално) PDF файлове като base64,
    извиква Claude API и връща генерирания текст.

    Тяло (JSON):
    {
        "prompt": "...",           # задължително — пълният prompt
        "files": [                 # опционално — прикачени документи
            {
                "name": "RS.pdf",
                "data": "<base64>",
                "media_type": "application/pdf"
            }
        ]
    }
    """
    if not ANTHROPIC_API_KEY:
        return jsonify({"error": "ANTHROPIC_API_KEY не е конфигуриран в Railway"}), 503

    body = request.get_json()
    if not body or "prompt" not in body:
        return jsonify({"error": "Липсва поле 'prompt'"}), 400

    prompt = body["prompt"]
    files  = body.get("files", [])

    try:
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

        # Изграждаме съдържанието на съобщението
        content = []

        # Добавяме прикачените PDF файлове
        for f in files:
            if f.get("media_type") == "application/pdf" and f.get("data"):
                content.append({
                    "type": "document",
                    "source": {
                        "type": "base64",
                        "media_type": "application/pdf",
                        "data": f["data"]
                    }
                })

        # Добавяме prompt-а като текст
        content.append({
            "type": "text",
            "text": prompt
        })

        response = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=4000,
            messages=[{"role": "user", "content": content}]
        )

        result_text = "".join(
            block.text for block in response.content
            if hasattr(block, "text")
        )

        return jsonify({
            "status": "ok",
            "result": result_text,
            "input_tokens":  response.usage.input_tokens,
            "output_tokens": response.usage.output_tokens,
        })

    except anthropic.APIError as e:
        return jsonify({"error": f"Claude API грешка: {str(e)}"}), 502
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/ai/chat", methods=["POST"])
def ai_chat():
    if not ANTHROPIC_API_KEY:
        return jsonify({"error": "ANTHROPIC_API_KEY не е конфигуриран"}), 503
    body = request.get_json()
    if not body or "messages" not in body:
        return jsonify({"error": "Липсва поле 'messages'"}), 400
    messages     = body["messages"]
    system_prompt = body.get("system", "")
    max_tokens   = int(body.get("max_tokens", 4000))
    try:
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        kwargs = dict(model="claude-sonnet-4-6", max_tokens=max_tokens, messages=messages)
        if system_prompt:
            kwargs["system"] = system_prompt
        response = client.messages.create(**kwargs)
        result_text = "".join(b.text for b in response.content if hasattr(b, "text"))
        return jsonify({
            "status": "ok",
            "result": result_text,
            "input_tokens":  response.usage.input_tokens,
            "output_tokens": response.usage.output_tokens,
        })
    except anthropic.APIError as e:
        return jsonify({"error": f"Claude API грешка: {str(e)}"}), 502
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/ai/extract-protokol2", methods=["POST"])
@require_auth
def ai_extract_protokol2():
    if not ANTHROPIC_API_KEY:
        return jsonify({"error": "ANTHROPIC_API_KEY не е конфигуриран"}), 503

    body = request.get_json()
    if not body:
        return jsonify({"error": "Липсва тяло на заявката"}), 400

    files     = body.get("files", [])
    pi        = str(body.get("pi", "unknown"))
    tenant_id = request.current_user.get("tenant_id")
    user_id   = request.current_user.get("sub")

    extract_prompt = """Ти си експерт по Bulgarian строителна документация. Анализирай приложените PDF файлове (Разрешение за строеж и/или одобрен архитектурен проект) и извлечи следните стойности.

Върни САМО валиден JSON обект — без markdown, без ```json, без обяснения. Ако дадена стойност не може да бъде намерена в документите, върни null за нея.

{
  "Opisanie_Ploshtadka": "<описание на сградата от архитектурните бележки: брой етажи, застроена площ, разстояния до регулация, конструктивна схема, височина>",
  "Kota_Izkop": "<кота изкоп в метри, напр. -2.50>",
  "Kota_Cokul": "<абсолютна кота ±0.00 в метри, напр. 142.50>",
  "Kota_Korniz": "<кота корниз/стреха в метри, напр. +9.80>",
  "Kota_Bilo": "<кота било/Ridge в метри, напр. +12.40>",
  "Reper_Nomer": "<номер на нивелетен репер, напр. РП-23 или Репер №5>",
  "Reper_Kota": "<кота на репера в метри, напр. 143.750>",
  "Darvesenost": null
}"""

    content = []
    for f in files:
        if f.get("media_type") == "application/pdf" and f.get("data"):
            content.append({
                "type": "document",
                "source": {
                    "type": "base64",
                    "media_type": "application/pdf",
                    "data": f["data"]
                }
            })
    content.append({"type": "text", "text": extract_prompt})

    if not any(c["type"] == "document" for c in content):
        return jsonify({"error": "Не са предоставени PDF файлове за анализ"}), 400

    try:
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        response = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=1000,
            messages=[{"role": "user", "content": content}]
        )

        raw = "".join(b.text for b in response.content if hasattr(b, "text")).strip()

        try:
            extracted = json.loads(raw)
        except json.JSONDecodeError:
            import re as _re
            m = _re.search(r'\{.*\}', raw, _re.DOTALL)
            if m:
                extracted = json.loads(m.group())
            else:
                return jsonify({"error": "Claude не върна валиден JSON", "raw": raw}), 502

        log_action("ai_extract_protokol2", user_id=user_id, tenant_id=tenant_id,
                   detail=f"PI={pi} tokens={response.usage.input_tokens}+{response.usage.output_tokens}")

        return jsonify({
            "status": "ok",
            "extracted": extracted,
            "input_tokens":  response.usage.input_tokens,
            "output_tokens": response.usage.output_tokens,
        })

    except anthropic.APIError as e:
        return jsonify({"error": f"Claude API грешка: {str(e)}"}), 502
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/ai/generate-akt15-sgrada", methods=["POST"])
@require_auth
def ai_generate_akt15_sgrada():
    """
    Акт 15 за сграда — приема prompt + файлове (PDF и docx) като base64,
    извиква Claude API и връща генерирания текст.
    Тяло (JSON): { "pi": "...", "prompt": "...", "files": [{name, data, media_type}] }
    """
    body = request.get_json()
    print(f"[DEBUG] Body keys: {list(body.keys()) if body else 'EMPTY'}", flush=True)

    manual = body.get("manual", {}) if body else {}
    apartments = body.get("apartments", []) if body else []

    # ── ТЕСТОВ РЕЖИМ — без извикване на Claude API ──
    if body and body.get("test_mode"):
        validation_result = validate_akt15(manual, apartments=apartments)
        return jsonify({
            "status": "ok",
            "result": "[ТЕСТОВ РЕЖИМ] Фиктивен текст на Акт 15 — реалното AI извикване е пропуснато.",
            "validation": validation_result,
            "input_tokens": 0,
            "output_tokens": 0,
        })
    # ── край на тестовия режим ──

    if not ANTHROPIC_API_KEY:
        return jsonify({"error": "ANTHROPIC_API_KEY не е конфигуриран"}), 503

    print(f"[DEBUG] Prompt present: {'prompt' in body if body else False}", flush=True)
    print(f"[DEBUG] Prompt length: {len(body.get('prompt', '')) if body else 0}", flush=True)
    if not body or "prompt" not in body:
        return jsonify({"error": "Липсва поле 'prompt'"}), 400

    pi        = str(body.get("pi", "unknown"))
    prompt    = body["prompt"]
    files     = body.get("files", [])
    tenant_id = request.current_user.get("tenant_id")
    user_id   = request.current_user.get("sub")

    try:
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

        content = []
        for f in files:
            mt = f.get("media_type", "")
            if not f.get("data"):
                continue
            if mt == "application/pdf":
                content.append({
                    "type": "document",
                    "source": {"type": "base64", "media_type": "application/pdf", "data": f["data"]}
                })
            elif "word" in mt or "openxmlformats" in mt:
                try:
                    raw = base64.b64decode(f["data"])
                    doc_obj = Document(io.BytesIO(raw))
                    docx_text = "\n".join(p.text for p in doc_obj.paragraphs if p.text.strip())
                    content.append({
                        "type": "text",
                        "text": f"=== СЪДЪРЖАНИЕ НА ДОКУМЕНТ: {f.get('name', 'проект.docx')} ===\n{docx_text}\n=== КРАЙ НА ДОКУМЕНТА ==="
                    })
                except Exception as docx_err:
                    print(f"[WARN] Неуспешно четене на .docx: {docx_err}", flush=True)
        content.append({"type": "text", "text": prompt})

        response = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=2000,
            messages=[{"role": "user", "content": content}]
        )

        result_text = "".join(
            block.text for block in response.content if hasattr(block, "text")
        )

        validation_result = validate_akt15(manual, apartments=apartments)
        print(f"[DEBUG] Validation result: {validation_result}", flush=True)

        log_action("generate_akt15_sgrada", user_id=user_id, tenant_id=tenant_id,
                   detail=f"PI={pi} tokens={response.usage.input_tokens}+{response.usage.output_tokens}")

        return jsonify({
            "status": "ok",
            "result": result_text,
            "validation": validation_result,
            "input_tokens":  response.usage.input_tokens,
            "output_tokens": response.usage.output_tokens,
        })

    except anthropic.APIError as e:
        return jsonify({"error": f"Claude API грешка: {str(e)}"}), 502
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ── AI Валидация ──────────────────────────────────────────────────────────────
VALIDATION_SYSTEM_PROMPT = """Ти си експерт по българското строително законодателство, специализиран в Наредба №3 от 2003 г. за съставяне на актове и протоколи по време на строителство.

Задачата ти е да валидираш строителни документи. Анализирай предоставения текст и върни САМО валиден JSON обект (без обяснения, без markdown, без ```json блокове) в следния формат:

{
  "status": "valid" | "warning" | "error",
  "score": <число 0-100>,
  "issues": [
    {"severity": "error"|"warning"|"info", "field": "<поле или секция>", "message": "<описание на проблема>"}
  ],
  "suggestions": [
    "<конкретна препоръка за подобрение>"
  ]
}

Правила за статус:
- "valid"   — score >= 80, няма грешки (error), може да има предупреждения
- "warning" — score 50-79, или има предупреждения без грешки
- "error"   — score < 50, или има поне една грешка (error)

Критерии за валидация:
1. ПЪЛНОТА — всички задължителни полета попълнени (страни, дати, описания, подписи)
2. ЛОГИЧЕСКА ПОСЛЕДОВАТЕЛНОСТ — съответства ли на изискванията на Наредба №3 за конкретния тип документ
3. РЕГУЛАТОРНО СЪОТВЕТСТВИЕ — правилни наименования, номерация, позовавания на нормативни актове
4. ФОРМАЛНИ ИЗИСКВАНИЯ — правилен формат на дати (дд.мм.гггг), номера на разрешения, ЕИК формат"""

DOC_TYPE_CONTEXT = {
    "protokol2":   "Протокол №2 за установяване годността за ползване на строежа (Приложение №2 към Наредба №3)",
    "protokol2p2": "Протокол №2, Част 2 — допълнителни констатации",
    "protokol2a":  "Протокол №2а по Наредба №3",
    "obrazec3":    "Образец №3 — констативен протокол",
    "akt5":        "Акт №5 за приемане на кофражи и армировка (Приложение №5 към Наредба №3)",
    "akt6":        "Акт №6 за приемане на направени изкопи (Приложение №6 към Наредба №3)",
    "akt7":        "Акт №7 за приемане на земна основа (Приложение №7 към Наредба №3)",
    "akt8":        "Акт №8 за приемане на изпълнена конструкция (Приложение №8 към Наредба №3)",
    "akt9":        "Акт №9 — Дневник за бетонови работи (Приложение №9 към Наредба №3)",
    "akt10":       "Акт №10 за приемане на изпълнена хидроизолация (Приложение №10 към Наредба №3)",
    "akt11":       "Акт №11 за приемане на изпълнена топлоизолация (Приложение №11 към Наредба №3)",
    "akt13":       "Акт №13 за приемане на конструкцията (Приложение №13 към Наредба №3)",
    "akt14":       "Акт №14 за приемане на конструкцията и частите на строежа (Приложение №14 към Наредба №3)",
    "zapovedna":   "Заповедна книга по чл.163 от ЗУТ",
    "akt15":       "Акт №15 за установяване на всички видове строителни и монтажни работи (Приложение №15 към Наредба №3)",
    "akt16":       "Акт №16 за установяване годността за ползване на строежа (Приложение №16 към Наредба №3)",
    "doklad":      "Окончателен доклад на консултанта по чл.168 от ЗУТ",
}


@app.route("/api/validate/<doc_type>", methods=["POST"])
@require_auth
def validate_document(doc_type):
    if not ANTHROPIC_API_KEY:
        return jsonify({"error": "ANTHROPIC_API_KEY не е конфигуриран в Railway"}), 503

    if doc_type not in DOC_TYPE_CONTEXT:
        return jsonify({"error": f"Непознат тип: {doc_type}. Позволени: {list(DOC_TYPE_CONTEXT.keys())}"}), 400

    body = request.get_json()
    if not body or not body.get("text"):
        return jsonify({"error": "Липсва поле 'text' с текста на документа"}), 400

    doc_text  = body["text"]
    doc_label = DOC_TYPE_CONTEXT[doc_type]
    tenant_id = request.current_user.get("tenant_id")
    user_id   = request.current_user.get("sub")

    user_prompt = f"""Тип документ: {doc_label}

Текст за валидация:
---
{doc_text}
---

Валидирай документа по четирите критерия и върни JSON."""

    try:
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        response = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=2000,
            system=VALIDATION_SYSTEM_PROMPT,
            messages=[{"role": "user", "content": user_prompt}]
        )

        raw = "".join(block.text for block in response.content if hasattr(block, "text")).strip()

        # Strip accidental markdown fences
        if raw.startswith("```"):
            raw = re.sub(r"^```[a-z]*\n?", "", raw)
            raw = re.sub(r"\n?```$", "", raw.rstrip())

        result = json.loads(raw)

        # Ensure required keys exist with safe defaults
        result.setdefault("status", "error")
        result.setdefault("score", 0)
        result.setdefault("issues", [])
        result.setdefault("suggestions", [])

        log_action("validate_doc", user_id=user_id, tenant_id=tenant_id,
                   detail={"doc_type": doc_type, "status": result["status"], "score": result["score"]})

        return jsonify({
            **result,
            "doc_type":      doc_type,
            "input_tokens":  response.usage.input_tokens,
            "output_tokens": response.usage.output_tokens,
        })

    except json.JSONDecodeError:
        return jsonify({"error": "Claude върна невалиден JSON", "raw": raw}), 502
    except anthropic.APIError as e:
        return jsonify({"error": f"Claude API грешка: {str(e)}"}), 502
    except Exception as e:
        return jsonify({"error": str(e)}), 500


init_db()

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    print(f"АКТ СИСТЕМ API v2.1 | порт {port} | шаблони: {TEMPLATES_FOLDER}")
    app.run(host="0.0.0.0", port=port)
