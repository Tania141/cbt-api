"""
cbt_docx.py — Document-building helpers for АКТ СИСТЕМ.
Extracted from api.py (refactor only, no behavior change).
"""
from datetime import datetime
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Name / date helpers ───────────────────────────────────────────────────────

_TITLES = {"инж", "арх", "проф", "д-р", "доц"}

def _split_title(s):
    """Отдели водеща титла от имената — титлата не бива да се брои като име."""
    parts = (s or "").strip().split()
    if parts and parts[0].lower().rstrip(".") in _TITLES:
        return parts[0], parts[1:]
    return "", parts

def two_names(s):
    title, parts = _split_title(s)
    return " ".join(([title] if title else []) + parts[:2])

def one_and_three(s):
    title, parts = _split_title(s)
    core = f"{parts[0]} {parts[2]}" if len(parts) >= 3 else " ".join(parts)
    return f"{title} {core}".strip() if core else ""

def with_title(title, name):
    """
    Слепи титла + име — единственото място, където титла се долепя.
    Ако името вече носи вградена титла (стари паспорти), не удвоява.
    """
    name = (name or "").strip()
    if not name:
        return ""
    title = (title or "").strip()
    if not title:
        return name
    if name.split()[0].lower().rstrip(".") in _TITLES:
        return name
    return f"{title} {name}"

def fmt_date(v):
    if not v: return ""
    v = v.strip()
    if "." in v: return v
    try: return datetime.fromisoformat(v[:10]).strftime("%d.%m.%Y")
    except: return v

def _clean_pred(s):
    """Махни водещо 'от ' — шаблонът вече долепя 'представлявано от'."""
    s = (s or "").strip()
    return s[3:].strip() if s.lower().startswith("от ") else s


# ── People extractors (used only by build_placeholders) ───────────────────────

def extract_employees(d):
    n = int(d.get("Служители_Брой", 0) or 0)
    return [{
        "name": d.get(f"Служител_{i}_Име",""),
        "specialization": d.get(f"Служител_{i}_Специализация",""),
        "title": d.get(f"Служител_{i}_Титла","инж.")
    } for i in range(1, n+1) if d.get(f"Служител_{i}_Име","")]

def extract_projectants(d):
    n = int(d.get("Проектанти_Брой", 0) or 0)
    return [{
        "name": d.get(f"Проектант_{i}_Име",""),
        "ppp": d.get(f"Проектант_{i}_ППП",""),
        "specialization": d.get(f"Проектант_{i}_Специализация",""),
        "title": d.get(f"Проектант_{i}_Титла","инж.")
    } for i in range(1, n+1) if d.get(f"Проектант_{i}_Име","")]

def by_spec(items, spec):
    return [e for e in items if e.get("specialization") == spec]


# ── Structural block builders ─────────────────────────────────────────────────

def extract_vazlogiteli(d):
    """Extract list of clients — same pattern as extract_employees/extract_projectants."""
    n = int(d.get("Възложители_Брой", 0) or 0)
    return [{
        "tip":            d.get(f"Възложител_{i}_Тип", "ФЛ"),
        "firma":          d.get(f"Възложител_{i}_Фирма", ""),
        "firma_title":    d.get(f"Възложител_{i}_Титла", ""),
        "eik":            d.get(f"Възложител_{i}_ЕИК", ""),
        "adres":          d.get(f"Възложител_{i}_Адрес", ""),
        "pred":           d.get(f"Възложител_{i}_Представител", ""),
        "pred_title":     d.get(f"Възложител_{i}_Представител_Титла", ""),
        "podpisva":       d.get(f"Възложител_{i}_Подписва", ""),
        "podpisva_title": d.get(f"Възложител_{i}_Подписва_Титла", ""),
        # липсващ ключ = подписва (обратна съвместимост със стария payload)
        "signs":          str(d.get(f"Възложител_{i}_Подписва_Да", "1")).strip() not in ("", "0", "false", "False", "не"),
    } for i in range(1, n+1) if d.get(f"Възложител_{i}_Фирма", "")]


def _single_vazlogitel_line(v):
    """Format one client entry for the header block."""
    if v["tip"] in ("Физическо лице", "ФЛ"):
        return with_title(v.get("firma_title"), v["firma"])
    parts = [v["firma"]]
    if v["eik"]:   parts.append(f"ЕИК {v['eik']}")
    if v["adres"]: parts.append(v["adres"])
    if v["pred"]:
        parts.append("представлявано от " + with_title(v.get("pred_title"), _clean_pred(v["pred"])))
    return ", ".join(parts)


def build_vazlogitel_block(d):
    """
    ГОРЕ блок ({{Възложител_Блок}}): номериран списък 1..N.
    Ако Възложители_Брой липсва → стар единичен път (обратна съвместимост).
    """
    vazlogiteli = extract_vazlogiteli(d)
    if vazlogiteli:
        if len(vazlogiteli) == 1:
            return _single_vazlogitel_line(vazlogiteli[0])
        lines = [f"{i+1}. {_single_vazlogitel_line(v)}" for i, v in enumerate(vazlogiteli)]
        return "\n".join(lines)

    # ── стар единичен път (обратна съвместимост) ──────────────────────────────
    tip   = d.get("Възложител_Тип", "Фирма")
    firma = d.get("Възложител_Фирма", "")
    adres = d.get("Възложител_Адрес", "")
    if tip in ("Физическо лице", "ФЛ"):
        return firma + (f", {adres}" if adres else "")
    eik  = d.get("Възложител_ЕИК", "")
    pred = d.get("Възложител_Представител", "")
    parts = [firma]
    if eik:   parts.append(f"ЕИК {eik}")
    if adres: parts.append(adres)
    if pred:  parts.append(f"представлявано от {_clean_pred(pred)}")
    return ", ".join(parts)


def build_vazlogitel_podpisva_block(d):
    """
    ДОЛУ/ПОДПИСИ блок ({{Възложител_Подписва_Блок}}).
    Ако Възложители_Брой липсва → стар единичен път (обратна съвместимост).
    """
    vazlogiteli = extract_vazlogiteli(d)
    if vazlogiteli:
        upalnom = d.get("Възложител_Упълномощен_Представител", "").strip()
        if upalnom:
            return upalnom
        signers = [v for v in vazlogiteli if v["signs"]] or vazlogiteli
        if len(signers) == 1:
            v = signers[0]
            return with_title(v.get("podpisva_title"), v["podpisva"]) or _single_vazlogitel_line(v)
        lines = []
        for i, v in enumerate(signers):
            podp = with_title(v.get("podpisva_title"), v["podpisva"]) or "………"
            lines.append(f"{i+1}. {_single_vazlogitel_line(v)} — подписва: {podp}")
        return "\n".join(lines)

    # ── стар единичен път (обратна съвместимост) ──────────────────────────────
    return d.get("Възложител_Подписва", "") or build_vazlogitel_block(d)


def build_vazlogitel_podpisva_redove(d):
    """
    Физически редове с точки за подписи — 'Б.: ......... (Иван Петров)'.
    Един ред на възложител. Замества {{Възложател_1и3}} в шаблоните.
    """
    def _signing_line(v):
        name = with_title(v.get("podpisva_title"), one_and_three(v["podpisva"])) if v["podpisva"] else ""
        return f"Б.:  ..............................   ({name})" if name else "Б.:  .............................."

    vazlogiteli = extract_vazlogiteli(d)
    if vazlogiteli:
        upalnom = d.get("Възложител_Упълномощен_Представител", "").strip()
        if upalnom:
            return f"Б.:  ..............................   ({one_and_three(upalnom)})"
        signers = [v for v in vazlogiteli if v["signs"]]
        if not signers:                       # никой избран → празен ред за подпис
            return "Б.:  .............................."
        return "\n".join(_signing_line(v) for v in signers)

    # ── стар единичен път ──────────────────────────────────────────────────────
    vaz_podpisva = d.get("Възложател_Подписва", "")
    vaz_tip   = d.get("Възложител_Тип", "Фирма")
    vaz_pr    = d.get("Възложител_Представител", "")
    vaz_firma = d.get("Възложител_Фирма", "")
    name = vaz_podpisva or (vaz_firma if vaz_tip in ("Физическо лице", "ФЛ") else vaz_pr)
    return f"Б.:  ..............................   ({one_and_three(name)})"


def build_projectants_list(projectants):
    lines = []
    for p in projectants:
        spec = p["specialization"]
        title = p["title"]
        name = p["name"]
        ppp = p["ppp"]
        kamara = "КАБ" if spec in ("Архитектура", "Паркоустройство и Благоустройство") else "КИИП"
        line = "Част " + spec + ": " + with_title(title, name)
        if ppp:
            line += ", рег. № " + ppp + " в " + kamara
        lines.append(line)
    return "\n".join(lines)

def build_employees_list(employees):
    return "\n".join(f"Част {e['specialization']}: {with_title(e['title'], e['name'])}" for e in employees)

def build_projectants_signatures(projectants):
    return "\n".join(
        f"Част {p['specialization']}: {with_title(p['title'], p['name'])} ….................................................................."
        for p in projectants
    )

def build_employees_signatures(employees):
    return "\n".join(
        f"Част {e['specialization']}: {with_title(e['title'], e['name'])} ….................................................................."
        for e in employees
    )

def _normalize_cokul(raw: str) -> str:
    """Strip accidental '+/-0,00=' prefix; return the tail VERBATIM (comma, spaces, ' м' intact)."""
    s = raw.strip()
    if "=" not in s:
        return s
    head, tail = s.split("=", 1)
    _pfx = head.replace(" ", "").replace("+/-", "±").replace(",", ".")
    if _pfx in ("±0.00", "±0", "0.00", "±0.0"):
        return tail.strip()
    return s


def build_placeholders(d):
    employees   = extract_employees(d)
    projectants = extract_projectants(d)

    # ── лица от списъци: титлата идва от собственото им поле ──────────────────
    def _person(items, spec, fallback_name=""):
        p = (by_spec(items, spec) or [{}])[0]
        return with_title(p.get("title", ""), p.get("name", "") or fallback_name)

    geo     = _person(employees,   "Геодезия", d.get("Геодезист", ""))
    sn_k    = _person(employees,   "Конструктивна")
    pj_k    = _person(projectants, "Конструктивна")
    pj_arch = _person(projectants, "Архитектура")
    specs   = "; ".join(f"{with_title(e['title'], e['name'])} ({e['specialization']})" for e in employees) \
              or d.get("Консултант_Управител","")
    vaz_tip = d.get("Възложител_Тип", "Фирма")
    # ── лица от свободен текст: титлата е в отделен ключ ──────────────────────
    upr     = with_title(d.get("Консултант_Управител_Титла",""), d.get("Консултант_Управител", ""))
    teh_ryk = with_title(d.get("Строител_ТехРък_Титла",""),      d.get("Строител_ТехРък", ""))
    str_upr = with_title(d.get("Строител_Управител_Титла",""),   d.get("Строител_Управител", ""))
    vaz_pr  = d.get("Възложител_Представител", "")
    # Възложител_1и3: prefer explicit signing person, fall back to ФЛ name or representative
    vaz_podpisva = d.get("Възложител_Подписва", "")
    vaz_name_for_1i3 = vaz_podpisva or (d.get("Възложител_Фирма", "") if vaz_tip in ("Физическо лице", "ФЛ") else vaz_pr)
    # при нов списъчен контракт вземи титлата от clients[0]
    _v0 = (extract_vazlogiteli(d) or [{}])[0]
    if _v0:
        _is_fl = _v0.get("tip") in ("Физическо лице", "ФЛ")
        vaz_name_for_1i3 = with_title(
            _v0.get("podpisva_title") if _v0.get("podpisva")
            else (_v0.get("firma_title") if _is_fl else _v0.get("pred_title")),
            _v0.get("podpisva") or (_v0.get("firma") if _is_fl else _clean_pred(_v0.get("pred", ""))),
        )

    return {
        "{{Строеж}}":                    d.get("Строеж",""),
        "{{Адрес}}":                     d.get("Адрес",""),
        "{{Консултант_Фирма}}":          d.get("Консултант_Фирма",""),
        "{{Консултант_ЕИК}}":            d.get("Консултант_ЕИК",""),
        "{{Консултант_Адрес}}":          d.get("Консултант_Адрес",""),
        "{{Консултант_Управител}}":      upr,
        "{{Консултант_Удостоверение}}":   d.get("Консултант_Удостоверение",""),
        "{{Управител_2имена}}":          two_names(upr),
        "{{Управител_1и3}}":             one_and_three(upr),
        "{{Строител_Фирма}}":            d.get("Строител_Фирма",""),
        "{{Строител_ЕИК}}":              d.get("Строител_ЕИК",""),
        "{{Строител_Адрес}}":            d.get("Строител_Адрес",""),
        "{{Строител_Управител}}":        str_upr,
        "{{Строител_Управител_2имена}}": two_names(str_upr),
        "{{Строител_Управител_1и3}}":    one_and_three(str_upr),
        "{{Строител_ТехРък}}":           teh_ryk,
        "{{ТехРък_2имена}}":             two_names(teh_ryk),
        "{{ТехРък_1и3}}":               one_and_three(teh_ryk),
        "{{tech_director}}":             teh_ryk,
        "{{Възложител_Тип}}":            vaz_tip,
        "{{Възложител_Фирма}}":          d.get("Възложител_Фирма",""),
        "{{Възложител_ЕИК}}":            d.get("Възложител_ЕИК","") if vaz_tip not in ("Физическо лице","ФЛ") else "",
        "{{Възложител_Адрес}}":          d.get("Възложител_Адрес",""),
        "{{Възложител_Представител}}":   vaz_pr if vaz_tip not in ("Физическо лице","ФЛ") else "",
        # ── при нов списъчен контракт: overwrite горните 4 от clients[0] ──────
        **({
            "{{Възложител_Тип}}":          _vl[0]["tip"],
            "{{Възложител_Фирма}}":        with_title(_vl[0]["firma_title"], _vl[0]["firma"])
                                           if _vl[0]["tip"] in ("Физическо лице","ФЛ") else _vl[0]["firma"],
            "{{Възложител_ЕИК}}":          _vl[0]["eik"],
            "{{Възложител_Адрес}}":        _vl[0]["adres"],
            "{{Възложител_Представител}}": with_title(_vl[0]["pred_title"], _clean_pred(_vl[0]["pred"])),
        } if (_vl := extract_vazlogiteli(d)) else {}),
        "{{Възложител_2имена}}":         two_names(vaz_name_for_1i3),
        "{{Възложител_1и3}}":           one_and_three(vaz_name_for_1i3),
        "{{Възложител_Блок}}":           build_vazlogitel_block(d),          # и (correct)
        "{{Възложател_Блок}}":           build_vazlogitel_block(d),          # а (legacy alias)
        "{{Възложител_Подписва_Блок}}":  build_vazlogitel_podpisva_block(d), # и
        "{{Възложател_Подписва_Блок}}":  build_vazlogitel_podpisva_block(d), # а (legacy alias)
        "{{Възложател_Подписва_Редове}}": build_vazlogitel_podpisva_redove(d),
        "{{РС_Номер}}":                  d.get("РС_Номер",""),
        "{{РС_Дата}}":                   fmt_date(d.get("РС_Дата","")),
        "{{РС_Издател}}":                d.get("РС_Издател",""),
        "{{РС_ВСила}}":                  fmt_date(d.get("РС_ВСила","")),
        "{{Геодезист}}":                 geo,
        "{{Геодезист_2имена}}":          two_names(geo),
        "{{Геодезист_1и3}}":            one_and_three(geo),
        "{{consultant_specialists}}":    specs,
        "{{sn_konstruktivna}}":          sn_k,
        "{{СН_Конструктивна}}":          sn_k,
        "{{СН_Архитектура}}":             _person(employees, "Архитектура"),
        "{{СН_Електро}}":                 _person(employees, "Електро"),
        "{{СН_ВиК}}":                     _person(employees, "ВиК"),
        "{{СН_Геодезия}}":                _person(employees, "Геодезия"),
        "{{СН_ПБ}}":                      _person(employees, "ПБ"),
        "{{СН_Пътна}}":                   _person(employees, "Пътна"),
        "{{СН_ОВК}}":                     _person(employees, "ОВК и ЕЕ"),
        "{{Конструктивна}}":              pj_k,
        "{{Конструктивна_1и3}}":          one_and_three(pj_k),
        "{{ПЖ_Конструктивна}}":           pj_k,
        "{{constructor_name}}":           pj_k,
        "{{ПЖ_Конструктивна_1и3}}":       one_and_three(pj_k),
        "{{ПЖ_Конструктивна _1и3}}":      one_and_three(pj_k),   # alias: space typo in template
        "{{ПЖ_Архитектура}}":             pj_arch,
        "{{ПЖ_Архитектура_1и3}}":         one_and_three(pj_arch),
        "{{ПЖ_Архитектура _1и3}}":        one_and_three(pj_arch), # alias: space typo in template
        "{{Вода}}":                      d.get("Вода", ""),
        "{{Канализация}}":               d.get("Канализация", ""),
        "{{Ел_Захранване}}":             d.get("Ел_Захранване", ""),
        "{{Проектанти_Списък}}":         build_projectants_list(projectants),
        "{{Консултанти_Списък}}":        build_employees_list(employees),
        "{{Проектанти_Подписи}}":        build_projectants_signatures(projectants),
        "{{Консултанти_Подписи}}":       build_employees_signatures(employees),
        "{{Opisanie_Ploshtadka}}":       d.get("Opisanie_Ploshtadka", ""),
        "{{Sastoyanie_Okolo}}":          d.get("Sastoyanie_Okolo", "пътните и тротоарни настилки по прилежащата улица са в добро състояние, съседните имоти няма да бъдат засягани от бъдещото строителство"),
        "{{Merki_PBZ}}":                 d.get("Merki_PBZ", "ще се осъществява от прилежащата улична мрежа съгласно съгласуван ПБЗ"),
        "{{Darvesenost}}":               d.get("Darvesenost", ""),
        "{{Kota_Izkop}}":               d.get("Kota_Izkop", ""),
        "{{Kota_Cokul}}":               _normalize_cokul(d.get("Kota_Cokul", "")),
        "{{Kota_Korniz}}":              d.get("Kota_Korniz", ""),
        "{{Kota_Bilo}}":                d.get("Kota_Bilo", ""),
        "{{Reper_Nomer}}":              d.get("Reper_Nomer", ""),
        "{{Reper_Kota}}":               d.get("Reper_Kota", ""),
        # Кирилични алиаси — същите стойности, различен правопис в Протокол 2
        "{{Кота_Изкоп}}":              d.get("Kota_Izkop", ""),
        "{{Кота_Цокъл}}":              _normalize_cokul(d.get("Kota_Cokul", "")),
        "{{Кота_Корниз}}":             d.get("Kota_Korniz", ""),
        "{{Кота_Bilo}}":               d.get("Kota_Bilo", ""),
        "{{Репер_Номер}}":             d.get("Reper_Nomer", ""),
        "{{Репер_Кота}}":              d.get("Reper_Kota", ""),
        # Описание на строителната площадка — ръчно поле
        "{{Описание_Сграда}}":         d.get("Opisanie_Sgrada", ""),
    }


# ── Template engine ───────────────────────────────────────────────────────────

def insert_paragraphs_after(para, lines, font_name="Times New Roman", font_size=12):
    from docx.shared import Pt
    ref = para._element
    parent = ref.getparent()
    idx = list(parent).index(ref)
    for i, line in enumerate(lines):
        new_p = OxmlElement("w:p")
        new_r = OxmlElement("w:r")
        new_rpr = OxmlElement("w:rPr")
        if para.runs:
            orig_rpr = para.runs[0]._r.find(qn("w:rPr"))
            if orig_rpr is not None:
                new_rpr = deepcopy(orig_rpr)
        new_r.append(new_rpr)
        new_t = OxmlElement("w:t")
        new_t.text = line
        new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        new_r.append(new_t)
        new_p.append(new_r)
        parent.insert(idx + 1 + i, new_p)

def replace_in_runs(para, replacements):
    full = "".join(r.text for r in para.runs)
    if not any(k in full for k in replacements):
        return False

    multiline_key = None
    multiline_val = None
    for k, v in replacements.items():
        if k in full and "\n" in v:
            multiline_key = k
            multiline_val = v
            break

    if multiline_key:
        lines = multiline_val.split("\n")
        first_line = full.replace(multiline_key, lines[0])
        for k, v in replacements.items():
            if k != multiline_key and "\n" not in v:
                first_line = first_line.replace(k, v)
        if para.runs:
            para.runs[0].text = first_line
            for r in para.runs[1:]:
                r.text = ""
        if len(lines) > 1:
            insert_paragraphs_after(para, lines[1:])
        return True

    new_text = full
    for k, v in replacements.items():
        new_text = new_text.replace(k, v)
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""
    return False

def fill_template(doc, replacements):
    paras = list(doc.paragraphs)
    for para in paras:
        replace_in_runs(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_runs(para, replacements)
    for section in doc.sections:
        for para in section.header.paragraphs:
            replace_in_runs(para, replacements)
        for para in section.footer.paragraphs:
            replace_in_runs(para, replacements)
    return doc
