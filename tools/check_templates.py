"""
check_templates.py — проверка на шаблоните срещу КОНВЕНЦИЯ.md

Пуска се от корена на cbt-api:
    python tools/check_templates.py

Две нива:
  ГРЕШКА   — нарушение на конвенцията; изходен код 1 (спира commit)
  ВНИМАНИЕ — за човешка преценка; не спира нищо
"""
import sys, os, re, glob

sys.stdout.reconfigure(encoding="utf-8")
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from cbt_docx import build_placeholders

ROOT    = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TPL_DIR = os.path.join(ROOT, "templates")

MARKER = re.compile(r"\{\{[^}]+\}\}")

# Маркер, който КОДЪТ нарочно записва: repl["{{Нещо}}"] = …
# Ако никой шаблон не го иска, стойността се изчислява и се изхвърля мълчаливо.
# Точно това стана на 02.09 със Заповедната книга: endpoint-ът пълнеше
# {{Заповедна_Номер}}, а шаблонът вече искаше {{ЗК_Номер}}. Номерът, въведен от
# оператора, не стигаше до документа. Като „сираче“ не се забеляза — сирачетата
# са близо шейсет и са безобидни, защото никой не се старае да ги попълни.
WRITTEN_MARKER = re.compile(r"""\w+\[\s*["'](\{\{[^}]+\}\})["']\s*\]\s*=""")
CODE_FILES = ("api.py", "cbt_docx.py")

# ── ГРЕШКИ: шаблонът дублира работата на кода ────────────────────────────────
# Блоковите маркери сами носят цялата фраза — пред тях не се пише нищо.
TITLE_BEFORE_MARKER = re.compile(r"(инж\.|арх\.|проф\.|доц\.|д-р)\s*\{\{")
PRED_BEFORE_BLOCK   = re.compile(r"представлявано от\s*\{\{[^}]*(Блок|Подписва)")
WRAPPED_REDOVE      = re.compile(r"[.\(]\s*\{\{[^}]*Подписва_Редове\}\}")
# Забележка: „Кота цокъл +/-0,00 = {{Kota_Cokul}}“ НЕ е дублиране — относителната
# кота на цокъла винаги е ±0,00, маркерът носи абсолютната (Вариант Б1, 29.07.2026).

# ── ВНИМАНИЕ: за човешко око ─────────────────────────────────────────────────
# Едно и също лице, изписано различно в РЕДОВЕ ОТ ЕДИН И СЪЩ ВИД.
# (дефектът от 13.08: „по част Архитектура" веднъж с пълно име, веднъж с _1и3)
SIBLING_LINE = re.compile(r"^\s*\d+\s*[.)]\s*(по част|\.{3,})")

NAME_FAMILIES = ["ПЖ_Архитектура", "ПЖ_Конструктивна", "Конструктивна",
                 "Геодезист", "Управител", "ТехРък", "Строител_Управител"]

LEGACY = {
    "Възложател_":               "правописен дублет с „а“ вместо „и“",
    "{{Кота_":                   "кирилски вариант на кота",
    "{{Репер_":                  "кирилски вариант на репер",
    " _1и3}}":                   "маркер с интервал (правописна грешка в шаблон)",
    "{{tech_director}}":         "латински вариант",
    "{{sn_konstruktivna}}":      "латински вариант",
    "{{consultant_specialists}}":"латински вариант",
}


def all_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts += [p.text for p in cell.paragraphs]
    for s in doc.sections:
        parts += [p.text for p in s.header.paragraphs]
        parts += [p.text for p in s.footer.paragraphs]
    return parts


def sibling_mismatch(lines):
    """Едно семейство имена, изписано и пълно, и съкратено, в еднотипни редове."""
    out = []
    for fam in NAME_FAMILIES:
        full, short = "{{%s}}" % fam, "{{%s_1и3}}" % fam
        sib_full = [l for l in lines if SIBLING_LINE.match(l) and full in l]
        sib_short = [l for l in lines if SIBLING_LINE.match(l) and short in l]
        if sib_full and sib_short:
            out.append(f"{fam}: пълно име и _1и3 в еднотипни редове "
                       f"({len(sib_full)} и {len(sib_short)} бр.)")
    return out


def main():
    known = set(build_placeholders({}).keys())
    files = sorted(glob.glob(os.path.join(TPL_DIR, "*.docx")))
    errors, warnings, legacy_hits, used = [], [], [], set()

    for path in files:
        name  = os.path.basename(path)
        lines = [l for l in all_text(Document(path))]
        blob  = "\n".join(lines)
        found = set(MARKER.findall(blob))
        used |= found

        for m in sorted(found - known):
            errors.append((name, "ВИСЯЩ МАРКЕР", f"{m} — няма ключ в build_placeholders"))

        for line in lines:
            for rx, what in [(TITLE_BEFORE_MARKER, "титла пред маркер"),
                             (PRED_BEFORE_BLOCK,   "„представлявано от“ пред блоков маркер"),
                             (WRAPPED_REDOVE,      "Подписва_Редове в скоби/точки")]:
                for mm in rx.finditer(line):
                    frag = line[max(0, mm.start() - 25):mm.end() + 30].strip()
                    errors.append((name, "ДУБЛИРАНЕ", f"{what}: …{frag}…"))

        for msg in sibling_mismatch(lines):
            warnings.append((name, "РАЗНОБОЙ", msg))

        for leg, why in LEGACY.items():
            if leg in blob:
                legacy_hits.append((name, leg, why))

    # Маркери, които кодът записва нарочно, а никой шаблон не иска.
    for fname in CODE_FILES:
        p = os.path.join(ROOT, fname)
        if not os.path.isfile(p):
            continue
        src = open(p, encoding="utf-8").read()
        for m in sorted(set(WRITTEN_MARKER.findall(src))):
            if m not in used:
                errors.append((fname, "МЪРТЪВ МАРКЕР",
                               f"{m} — кодът го записва, но никой шаблон не го иска; "
                               f"стойността се изхвърля мълчаливо"))

    orphans = sorted(known - used)

    print(f"Проверени шаблони: {len(files)}\n")

    if errors:
        print("── ГРЕШКИ ────────────────────────────────────────────────")
        cur = None
        for f, kind, msg in errors:
            if f != cur: print(f"\n  {f}"); cur = f
            print(f"     [{kind}] {msg[:120]}")
        print()
    else:
        print("✅ ГРЕШКИ: няма. Шаблоните спазват конвенцията.\n")

    if warnings:
        print("── ВНИМАНИЕ (за преценка, не спира commit) ───────────────")
        for f, kind, msg in warnings:
            print(f"   {f:40} [{kind}] {msg}")
        print()

    if legacy_hits:
        print("── НАСЛЕДЕНИ МАРКЕРИ ─────────────────────────────────────")
        for f, leg, why in legacy_hits:
            print(f"   {f:40} {leg:28} {why}")
        print()

    if orphans:
        print(f"── СИРАЧЕТА: {len(orphans)} ключа в кода без маркер в шаблон ──")
        print("   " + ", ".join(orphans))
        print("   (не е грешка — маркерът е на разположение, ако потрябва)\n")

    return 1 if errors else 0


if __name__ == "__main__":
    sys.exit(main())
