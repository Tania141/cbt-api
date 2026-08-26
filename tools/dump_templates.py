"""
dump_templates.py — генерира .md огледало на всеки .docx шаблон.

    python tools/dump_templates.py

.docx е единственият източник; .md се презаписва и служи само за да
се вижда в git какво точно е променено в шаблона.
"""
import sys, os, glob

sys.stdout.reconfigure(encoding="utf-8")
from docx import Document

ROOT    = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TPL_DIR = os.path.join(ROOT, "templates")

HEADER = ("<!-- Генериран автоматично от {name} чрез tools/dump_templates.py.\n"
          "     НЕ редактирай този файл — промените се правят в .docx. -->\n\n")


def block_items(doc):
    """Параграфи и таблици в реда, в който стоят в документа."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc)


def render(path):
    doc = Document(path)
    out, blanks = [], 0
    for item in block_items(doc):
        if hasattr(item, "rows"):                      # таблица
            if blanks: out.append("")
            blanks = 0
            for row in item.rows:
                cells = [" ".join(p.text.split()) for p in
                         (c.paragraphs[0] for c in row.cells)] if row.cells else []
                cells = [" ".join(c.text.split()) or "" for c in row.cells]
                out.append("| " + " | ".join(cells) + " |")
            out.append("")
        else:                                          # параграф
            text = item.text.rstrip()
            if not text.strip():
                blanks += 1
                continue
            if blanks:
                out.append("")
                blanks = 0
            style = ((item.style.name if item.style is not None else "") or "").lower()
            if style.startswith("heading"):
                out.append("### " + text.strip())
            else:
                out.append(text)
    # колонтитули
    heads = [p.text.strip() for s in doc.sections for p in s.header.paragraphs if p.text.strip()]
    foots = [p.text.strip() for s in doc.sections for p in s.footer.paragraphs if p.text.strip()]
    if heads:
        out += ["", "<!-- колонтитул горе -->"] + heads
    if foots:
        out += ["", "<!-- колонтитул долу -->"] + foots
    return "\n".join(out).rstrip() + "\n"


def main():
    files = sorted(glob.glob(os.path.join(TPL_DIR, "*.docx")))
    changed = 0
    for path in files:
        name = os.path.basename(path)
        md_path = os.path.splitext(path)[0] + ".md"
        new = HEADER.format(name=name) + render(path)
        old = open(md_path, encoding="utf-8").read() if os.path.exists(md_path) else None
        if new != old:
            open(md_path, "w", encoding="utf-8", newline="\n").write(new)
            changed += 1
            print(f"  обновен  {os.path.basename(md_path)}")
    print(f"\nШаблони: {len(files)} · обновени .md: {changed}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
